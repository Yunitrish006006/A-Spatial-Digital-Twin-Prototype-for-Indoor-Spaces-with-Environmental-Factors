from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

OUT_DIR = Path('/mnt/data')
base_script_path = OUT_DIR / 'lagrange_error_constant_speaker_script_google_slides.txt'
base = base_script_path.read_text(encoding='utf-8')

# Remove generic lines that are not necessary in the report body? Keep source and usage at top.
# Build glossary supplements keyed by slide.
glossary = {
1: [
    ("線性 Lagrange 插值", "用一個一次函數／平面去通過元素節點上的函數值。在三角形上，三個頂點值決定一個平面。", "教授若問『為什麼叫 Lagrange？』可答：因為它用節點值建構插值基底，滿足在本節點為 1、其他節點為 0 的 nodal basis 性質。"),
    ("誤差常數 C^L(K)", "它不是某次計算的誤差，而是把所有可能函數中最壞插值誤差統一控制住的比例常數。", "可答：C^L(K) 越小，代表同樣曲率下，這個三角形元素上的最大點誤差越容易被控制。"),
    ("最大範數／L∞ 範數", "在整個區域中取絕對值最大的那個點；不是平均，而是最壞點。", "可答：L∞ 適合逐點保證，因為它要求每一點都不能超過這個最大誤差界。"),
],
2: [
    ("Top-down 報告法", "先講問題、結論與直覺，再逐步展開技術細節。", "可答：因為這篇論文公式較多，先建立主線可避免聽眾一開始卡在符號。"),
    ("有限元素法 FEM", "把連續區域切成小元素，並用有限個基底函數近似原本無限維的函數問題。", "可答：本文使用 FEM 不是為了求 PDE 解，而是為了把誤差常數的最佳化問題離散化。"),
    ("Fujino–Morley 空間", "一種非共形有限元素空間，適合處理含二階導數的 H² 型問題。", "可答：它的正交分解讓 H² 能量可以寫成矩陣二次型，這是後面算 λ 的基礎。"),
    ("Bernstein 多項式", "一種具有凸包性質的多項式表示法，可用係數控制函數最大值。", "可答：本文用它把『無限多點的最大值限制』轉成『有限多個係數的最大值限制』。"),
],
3: [
    ("插值 interpolation", "已知節點值後，用較簡單的函數去補出區域內其他位置的近似值。", "可答：插值要求在節點上完全相同；近似則不一定需要通過所有節點。"),
    ("最大點誤差", "函數 u 與插值函數 Π^L u 在區域 K 中差距最大的點。", "可答：它比平均誤差保守，因為只要某一點誤差大，L∞ 範數就會反映出來。"),
    ("曲率直覺", "線性平面無法表現彎曲；函數越彎，線性插值通常越容易產生誤差。", "可答：論文用 H² 半範數量化這種彎曲程度。"),
    ("三角形元素 K", "有限元素法中局部計算的基本區塊；本論文所有誤差常數都先在單一三角形上估計。", "可答：先把單一元素的常數掌握住，才有機會推到整體網格誤差估計。"),
],
4: [
    ("問題層／方法層／結果層", "問題層定義要估什麼；方法層說明怎麼估；結果層給出數值上下界。", "可答：這三層能把論文貢獻和單純數值實驗區分開。"),
    ("嚴格上下界", "下界表示真實最佳常數至少多大；上界表示真實最佳常數不超過多大。", "可答：上下界越接近，代表我們越精確知道最佳常數的位置。"),
    ("相對誤差小於 1%", "上下界之間的距離相對於常數大小很小，說明估計很銳利。", "可答：本文不只是可計算，而是能用很窄的範圍夾住最佳常數。"),
    ("H² 半範數", "只量測二階導數大小，不含函數本身與一階導數的完整大小。", "可答：在插值誤差估計中，二階導數代表函數偏離線性的程度，所以用它控制線性插值誤差很自然。"),
],
5: [
    ("範數 norm", "衡量函數或向量大小的工具。不同範數看見的『大小』不同，例如平均大小、能量大小或最大值。", "可答：L² 範數像均方平均；L∞ 範數看最大點；H² 半範數看二階導數能量。"),
    ("L∞ 範數", "\u201c∞\u201d 代表取極限意義下最大的大小；對函數來說就是 sup |f(x)|。", "可答：它比 L² 更適合保證每一點的誤差上限。"),
    ("H² 半範數 |u|₂,K", "在區域 K 上把 u 的所有二階偏導數用 L² 方式量起來。", "可答：它是曲率能量，不是函數高度本身；線性函數的二階導數為零。"),
    ("λ 與 C^L 的關係", "論文把 C^L 的估計轉成 λ 的估計，核心關係是 C^L(K)=1/√λ(K)。", "可答：因為 λ 是『能量／最大誤差平方』的最小值；最小能量越大，代表要達到單位最大誤差越難，所以誤差常數越小。"),
    ("最壞比例", "在所有非零 admissible 函數中取最大誤差與曲率尺度的最大比值。", "可答：這就是最佳常數的本質；它保證對所有函數都成立。"),
],
6: [
    ("一維 Lagrange 插值", "在區間端點固定函數值，用一條直線連接兩端點。", "可答：u(x)=x² 時插值線是 x，因為端點值為 0 與 1。"),
    ("二階導數", "描述函數彎曲程度；x² 的二階導數是常數 2。", "可答：線性插值誤差與二階導數相關，因為線性函數本身沒有曲率。"),
    ("1/8 常數", "在一維最大範數、二階導數最大值控制下的經典線性插值誤差常數。", "可答：對 u=x²，最大誤差是 1/4，二階導數最大值是 2，所以比例為 1/8。"),
    ("為什麼用一維暖身", "一維情況能用簡單函數直接算出常數，幫助理解二維問題的結構。", "可答：二維三角形是同一概念的推廣，但要處理方向、形狀與無限多點最大值。"),
],
7: [
    ("無限維函數空間", "函數的自由度不是有限個數字，而是整個連續區域上的函數形狀。", "可答：直接在 H²(K) 中找最壞函數不可行，所以要離散化。"),
    ("離散化", "把連續問題轉成有限維代數問題，例如用有限個基底函數與係數表示函數。", "可答：離散化後，函數 u 變成係數向量 x。"),
    ("基底函數 φ_i", "有限元素空間中用來組合近似函數的基本函數。", "可答：任何離散函數可以寫成 Σ x_i φ_i。"),
    ("矩陣二次型 xᵀAx", "用矩陣 A 表示能量；xᵀAx 是係數向量 x 對應的 H² 能量。", "可答：A 類似剛性矩陣，記錄基底函數二階導數之間的內積。"),
    ("非共形元素", "有限元素函數不一定完全屬於原本要求的連續函數空間，但透過特殊自由度與分析仍可給出有效估計。", "可答：Fujino–Morley 的價值在於適合 H² 型估計與正交分解。"),
],
8: [
    ("凸包 convex hull", "包含一組點的最小凸集合；直觀上像用橡皮筋套住所有控制點形成的範圍。", "可答：若函數值落在控制係數的凸包內，就可用控制係數的最大值界住函數最大值。"),
    ("凸包性質", "Bernstein 基底非負且總和為 1，所以多項式值是控制係數的加權平均。", "可答：加權平均不會超過最大係數，也不會低於最小係數。"),
    ("Bernstein 控制係數", "多項式在 Bernstein 表示中的係數；它們不是一定等於函數在某些點的值，但可以控制函數值範圍。", "可答：本文利用係數最大值給 L∞ 範數上界。"),
    ("放寬 relaxation", "把原本難解的限制換成較容易處理但仍保守的限制。", "可答：這會得到可驗證的界；雖然可能不完全等價，但數值結果顯示界很銳利。"),
    ("矩陣 B", "把有限元素係數 x 轉換成 Bernstein 控制係數的矩陣。", "可答：因此 ||Bx||∞ 就能控制函數最大範數。"),
],
9: [
    ("最佳化問題 optimization problem", "在限制條件下最大化或最小化某個目標函數。本文是找滿足最大範數條件時最小的 H² 能量。", "可答：這個最小能量就是 λ 的離散版本。"),
    ("限制條件 constraint", "最佳化時必須滿足的條件；這裡是函數最大範數至少為 1。", "可答：把最大範數固定住，再找最小能量，才能得到最壞比例的倒數。"),
    ("λ_{h,B}", "h 表示網格離散尺度，B 表示用 Bernstein 放寬後得到的 λ 估計。", "可答：它用於推導 C^L 的上界。"),
    ("Cholesky 分解", "把正定矩陣分解為三角矩陣乘積的線性代數工具。", "可答：論文用它簡化放寬後問題的計算，讓 λ_{h,B} 可由某些行向量範數求得。"),
    ("界的方向", "λ 越大，C=1/√λ 越小；因此 λ 的下界會給 C 的上界。", "可答：因為倒平方根是遞減函數，這點在回答上下界問題時很重要。"),
],
10: [
    ("單位等腰直角三角形 K_{1,π/2}", "兩股長度相同且夾角為 90 度的基準三角形。", "可答：它常用作參考元素，且本文對它給出代表性結果。"),
    ("λ_{h,B}=5.7812", "論文 Table 2 中對此三角形與特定網格計算得到的 Bernstein 放寬 λ 值。", "可答：代入 C^L≤1/√λ_{h,B} 得上界 0.41596。"),
    ("倒平方根計算", "√5.7812≈2.4044，1/2.4044≈0.41596。", "可答：這只是最後換算；真正困難在於嚴格得到 λ_{h,B}。"),
    ("上下界寬度", "0.41596−0.40432≈0.01164。", "可答：約 2.8% 的絕對相對範圍；說明估計已相當集中。"),
    ("嚴格估計 vs 近似數值", "嚴格估計會控制捨入與演算法造成的不確定性，不只是浮點數輸出。", "可答：論文也提到使用區間算術來處理嚴格性問題。"),
],
11: [
    ("π/6、π/4、π/3、π/2", "代表不同三角形角度；常用弧度表示，分別是 30、45、60、90 度。", "可答：表格比較的是不同形狀下 C^L 的上下界。"),
    ("元素形狀 effect", "三角形角度與邊長比例會影響誤差常數。", "可答：形狀越退化，誤差常數通常越難控制。"),
    ("銳利 sharp", "上下界非常接近，表示估計接近真實最佳常數。", "可答：sharp 不只是數字小，而是界限夾得緊。"),
    ("Table 2 解讀", "重點不是背每個數字，而是看常數如何隨形狀變化，以及上下界是否接近。", "可答：π/3 附近數值較小，π/2 例子約 0.41，顯示形狀有明顯影響。"),
],
12: [
    ("形狀正則 shape regular", "一族三角形的最小角不趨近 0，最大角不趨近 π，元素不過度扁長。", "可答：shape regular 可防止誤差常數失控。"),
    ("退化三角形", "三角形接近一條線段，或某邊長趨近 0，導致幾何品質變差。", "可答：論文指出這時 C^L 可能趨近無窮。"),
    ("最大角接近 π", "三角形幾乎被拉平成線段，局部幾何變得病態。", "可答：在這種情況，即使函數曲率尺度相同，插值最大誤差也可能被放大。"),
    ("網格品質", "有限元素網格中元素的角度與長寬比品質。", "可答：這不是視覺問題，而是直接影響誤差界常數。"),
],
13: [
    ("收斂階數 vs 明確常數", "收斂階數告訴誤差隨 h 變小的速度；明確常數告訴實際界線前面的倍數。", "可答：嚴格保證需要階數與常數兩者。"),
    ("逐點誤差估計", "對每一個空間位置給出誤差界，而不是只保證平均誤差。", "可答：L∞ 範數正是逐點誤差控制的自然工具。"),
    ("PDE 有限元素解", "用 FEM 近似偏微分方程的解。", "可答：本文常數可作為後續最大範數誤差分析的局部插值估計基礎。"),
    ("可移植性", "一種方法可延伸到其他插值、元素或誤差估計問題的可能性。", "可答：Bernstein 處理最大範數的想法可能可用於其他高階元素。"),
],
14: [
    ("三句話總結", "用問題、方法、結果收束，不再新增技術細節。", "可答：口試或報告最後要讓教授記得主貢獻，不必重開推導。"),
    ("Fujino–Morley + Bernstein 的分工", "前者處理 H² 能量，後者處理 L∞ 最大值。", "可答：它們剛好對應核心不等式右邊與左邊。"),
    ("最佳常數", "讓不等式對所有函數都成立時可取的最小常數。", "可答：本文不是只找某個可用常數，而是逼近最佳常數。"),
    ("研究價值", "把抽象最壞誤差常數轉成可算且可驗證的上下界。", "可答：這是理論誤差分析與實際數值計算之間的橋接。"),
],
15: [
    ("Q&A 策略", "先辨認教授問的是定義、方法、數值，還是應用，再用一句核心回答加一個例子。", "可答：例如問凸包，就回答『Bernstein 基底非負且總和為一，所以函數值被控制係數界住』。"),
    ("如果被問 λ 為什麼要用", "因為直接求 C^L 是最大比值問題；轉成 λ 後可變成最小能量問題，更適合矩陣最佳化。", "可答：最後再用 C^L=1/√λ 轉回來。"),
    ("如果被問放寬是否不精確", "是放寬，所以理論上可能保守；但上下界數值接近，表示在本文例子中相當銳利。", "可答：這也是 Table 2 的重要意義。"),
    ("如果被問和 FEM 解 PDE 的關係", "插值誤差估計通常是 FEM 解誤差估計的重要組成部分；本文提供最大範數下更明確的局部常數。", "可答：它本身不直接解 PDE，但支援後續 PDE 最大誤差分析。"),
]
}

# Parse base slides into sections
parts = re.split(r'(?=Slide \d+｜)', base)
header = parts[0].strip()
slides = []
for part in parts[1:]:
    m = re.match(r'Slide (\d+)｜([^\n]+)\n(.*)', part, flags=re.S)
    if not m:
        continue
    num = int(m.group(1)); title = m.group(2).strip(); body = m.group(3).strip()
    slides.append((num, title, body))

out_lines = []
out_lines.append('線性 Lagrange 插值誤差常數估計｜逐頁口述講稿＋名詞補充版')
out_lines.append('')
out_lines.append('來源論文：Shirley Mae Galindo, Koichiro Ike, and Xuefeng Liu, “Error-constant estimation under the maximum norm for linear Lagrange interpolation,” Journal of Inequalities and Applications, 2022:109.')
out_lines.append('使用方式：每頁先照「口述講稿」說明；若教授追問，直接參考該頁後方「本頁名詞補充」。')
out_lines.append('')
out_lines.append('總答題原則：先用一句白話定義回答，再補一個與本頁內容相連的數學意義；不要一開始就展開完整推導。')
out_lines.append('')
for num,title,body in slides:
    out_lines.append(f'Slide {num}｜{title}')
    out_lines.append('')
    out_lines.append('【口述講稿】')
    out_lines.append(body)
    out_lines.append('')
    out_lines.append('【本頁名詞補充：教授可能問法與回答】')
    for term, definition, answer in glossary.get(num, []):
        out_lines.append(f'• {term}')
        out_lines.append(f'  - 白話定義：{definition}')
        out_lines.append(f'  - 被問時可答：{answer}')
    out_lines.append('')
    out_lines.append('—' * 28)
    out_lines.append('')

updated_text = '\n'.join(out_lines)
txt_path = OUT_DIR / 'lagrange_speaker_script_with_terms.txt'
txt_path.write_text(updated_text, encoding='utf-8')

# Create DOCX
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.65)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

styles = doc.styles
styles['Normal'].font.name = 'Noto Sans CJK TC'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK TC')
styles['Normal'].font.size = Pt(10.5)
styles['Normal'].paragraph_format.line_spacing = 1.15
styles['Normal'].paragraph_format.space_after = Pt(4)

for sname, size, color, bold in [('Title', 20, '1F2937', True), ('Heading 1', 16, '1F2937', True), ('Heading 2', 12.5, '374151', True)]:
    st = styles[sname]
    st.font.name = 'Noto Sans CJK TC'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK TC')
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = RGBColor.from_string(color)

# helper to shade paragraph
def shade_paragraph(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def add_small_box(text, fill='F3F4F6'):
    p = doc.add_paragraph()
    shade_paragraph(p, fill)
    r = p.add_run(text)
    r.font.name = 'Noto Sans CJK TC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK TC')
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(55,65,81)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_after = Pt(8)
    return p

p = doc.add_paragraph()
p.style = styles['Title']
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('線性 Lagrange 插值誤差常數估計')
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('逐頁口述講稿＋名詞補充版')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(75,85,99)
add_small_box('使用方式：每頁先照「口述講稿」說明；若教授追問，直接看該頁後方「本頁名詞補充」。')
add_small_box('答題原則：先用一句白話定義回答，再補一個與本頁內容相連的數學意義；不要一開始就展開完整推導。', fill='EEF2FF')

for num,title,body in slides:
    h = doc.add_paragraph(style='Heading 1')
    h.add_run(f'Slide {num}｜{title}')
    h2 = doc.add_paragraph(style='Heading 2')
    h2.add_run('口述講稿')
    for para in body.split('\n\n'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.18)
        p.add_run(para.replace('\n', ' '))
    h2 = doc.add_paragraph(style='Heading 2')
    h2.add_run('本頁名詞補充：教授可能問法與回答')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '名詞'
    hdr[1].text = '白話定義'
    hdr[2].text = '被問時可答'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Noto Sans CJK TC'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK TC')
                r.font.size = Pt(9.5)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E5E7EB')
        tcPr.append(shd)
    for term, definition, answer in glossary.get(num, []):
        row = table.add_row().cells
        row[0].text = term
        row[1].text = definition
        row[2].text = answer
        for cell in row:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = 'Noto Sans CJK TC'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK TC')
                    r.font.size = Pt(8.8)
    doc.add_paragraph('')

# Add compact global glossary at end for fast review
h = doc.add_paragraph(style='Heading 1')
h.add_run('最後快速複習：高頻名詞一句話版')
quick_terms = [
    ('範數', '衡量函數或向量大小的方法；L∞ 看最大值，L² 看平方平均，H² 半範數看二階導數能量。'),
    ('凸包', '包住一組點的最小凸集合；Bernstein 表示中函數值可被控制係數的凸包界住。'),
    ('FEM', '用有限個基底函數把連續函數問題轉成矩陣問題。'),
    ('Fujino–Morley', '適合 H² 型問題的非共形有限元素空間，可把二階導數能量轉成 xᵀAx。'),
    ('Bernstein', '用非負且總和為一的基底表示多項式，因此能用係數控制最大值。'),
    ('λ', '能量比例的最小值；C^L=1/√λ。'),
    ('shape regular', '三角形不過度扁長、角度不退化，讓誤差常數保持有界。'),
]
for term, expl in quick_terms:
    p = doc.add_paragraph(style=None)
    r = p.add_run(term + '：')
    r.bold = True
    p.add_run(expl)

docx_path = OUT_DIR / 'lagrange_speaker_script_with_terms.docx'
doc.save(docx_path)
print(txt_path)
print(docx_path)
